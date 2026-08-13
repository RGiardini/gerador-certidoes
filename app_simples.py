import streamlit as st
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from io import BytesIO
import datetime
import locale

# Tenta configurar para português
try:
    locale.setlocale(locale.LC_TIME, 'pt_BR.utf8')
except:
    pass

st.set_page_config(page_title="Certidão Negativa", layout="centered")

# --- ESTILO COMPACTO (CSS) ---
st.markdown("""
    <style>
    .block-container { padding-top: 1rem; padding-bottom: 1rem; }
    h1 { font-size: 22px; text-align: center; margin-bottom: 0; padding-bottom: 0;}
    div[role="radiogroup"] { margin-top: -10px; }
    </style>
""", unsafe_allow_html=True)

st.title("Certidão Negativa")

# --- CABEÇALHO DO FORMULÁRIO ---
col_mandado, col_proc, col_ano = st.columns([1, 2, 1])
with col_mandado:
    mandado = st.text_input("Mandado nº:", placeholder="Ex: 01")
with col_proc:
    processo = st.text_input("Informe o Processo:", placeholder="Ex: 4400281-16")
with col_ano:
    ano = st.text_input("Ano:", placeholder="Ex: 2026")

# --- DATAS E HORÁRIOS ---
st.write("**Informe os Dias e Horários:**")
c_d1, c_d2, c_d3 = st.columns(3)
with c_d1:
    d1 = st.text_input("Dia 1", placeholder="Ex: 08/08")
    h1 = st.text_input("Hora 1", placeholder="Ex: 14:55hs")
with c_d2:
    d2 = st.text_input("Dia 2", placeholder="Ex: 11/08")
    h2 = st.text_input("Hora 2", placeholder="Ex: 16:58hs")
with c_d3:
    d3 = st.text_input("Dia 3", placeholder="Ex: 12/08")
    h3 = st.text_input("Hora 3", placeholder="Ex: 11:15hs")

# --- DESFECHO PRINCIPAL ---
situacao = st.radio(
    "Situação Principal:", 
    ["Local Fechado", "Pessoa Não Encontrada", "Não Localizei a Pessoa"],
    index=None, horizontal=True
)

# --- INFORMANTE ---
st.divider()
c_inf1, c_inf2 = st.columns([1, 2])
with c_inf1:
    obteve_inf = st.radio("Obteve Informações?", ["Sim", "Não", "NQI"], index=None, horizontal=True)
with c_inf2:
    nome_inf = st.text_input("Nome do Informante:", disabled=(obteve_inf != "Sim"))

# --- MOTIVOS E PARADEIRO ---
st.write("**Detalhes das Informações Obtidas:**")
c_m1, c_m2 = st.columns(2)
with c_m1:
    motivo = st.radio(
        "Motivo:", 
        ["Mudou-se", "Não Reside no Local", "Não fica ali", "Não trabalha ali", "Falecido"], 
        index=None
    )
with c_m2:
    nao_sabe = st.radio(
        "O que não sabe?", 
        ["Não Conhece ele", "Não sabe informar", "Não sabe seu endereço"], 
        index=None
    )
    paradeiro = st.radio(
        "Paradeiro:", 
        ["Não sabe o paradeiro", "Incerto e Não Sabido"], 
        index=None
    )

# --- CONDIÇÕES EXTRAS ---
st.divider()
condicao = st.radio(
    "Condições do Local:", 
    ["Local Perigoso", "Medo Processo", "Zona Rural", "Blocos", "Chuva"], 
    index=None, horizontal=True
)

# --- TEXTOS LIVRES ---
observacoes = st.text_area("Observações Extras:", height=68)
c_end, c_pes = st.columns(2)
with c_end:
    endereco = st.text_input("Endereço (opcional):", placeholder="Se vazio, usará 'informado no mesmo'")
with c_pes:
    pessoa = st.text_input("Pessoa procurada:", placeholder="Nome do destinatário...")

st.divider()

# ==========================================
# LÓGICA DE MONTAGEM DO TEXTO 
# ==========================================
if st.button("Salvar / Gerar DOCX", type="primary", use_container_width=True):
    
    dias_validos = [d for d in [d1, d2, d3] if d]
    horas_validas = [h for h in [h1, h2, h3] if h]
    
    texto_data_hora = ""
    if len(dias_validos) == 1:
        texto_data_hora = f", onde às {horas_validas[0]}, do dia {dias_validos[0]},"
    elif len(dias_validos) > 1:
        str_horas = ", ".join(horas_validas[:-1]) + f" e {horas_validas[-1]}"
        str_dias = ", ".join(dias_validos[:-1]) + f" e {dias_validos[-1]}"
        texto_data_hora = f", onde às {str_horas}, dos dias {str_dias},"

    txt_endereco = f"à {endereco}" if endereco else "ao endereço informado no mesmo"
    txt_pessoa = f" a pessoa, Sr(a). {pessoa}" if pessoa else "a pessoa referida no mandado"

    txt_situacao = ""
    if situacao == "Local Fechado":
        txt_situacao = "porque o local foi encontrado fechado e mesmo após chamar várias vezes, ninguém atendeu. "
    elif situacao == "Pessoa Não Encontrada":
        txt_situacao = "porque não a encontrei no local. "
    elif situacao == "Não Localizei a Pessoa":
        txt_situacao = "porque não a localizei. "

    paragrafo_unico = (
        f"Certifico e dou fé que, em cumprimento ao mandado anexo, dirigi-me {txt_endereco}{texto_data_hora} "
        f"e, deixei de citar/intimar/notificar {txt_pessoa}, {txt_situacao}"
    )

    if obteve_inf == "Sim":
        paragrafo_unico += f"Conforme informações obtidas no local com Sr.(a) {nome_inf}, informou que, "
    elif obteve_inf == "Não":
        paragrafo_unico += "Procurei obter informações junto aos moradores vizinhos locais, e não obtive êxito, uma vez que ninguém forneceu informações. "
    elif obteve_inf == "NQI":
        paragrafo_unico += "Conforme informações prestadas pelo seu vizinho(a), que não quis se identificar, este afirmou que "

    if obteve_inf in ["Sim", "NQI"]:
        if motivo == "Mudou-se":
            paragrafo_unico += "a pessoa procurada não reside mais no local, tendo se mudado sem deixar meios para contato; "
        elif motivo == "Não Reside no Local":
            paragrafo_unico += "a pessoa procurada não reside no local referido; "
        elif motivo == "Não fica ali":
            paragrafo_unico += "a pessoa procurada reside no local, mas quase não fica no mesmo, onde nos dias e horários acima não foi localizada; "
        elif motivo == "Não trabalha ali":
            paragrafo_unico += "a pessoa procurada não trabalha no local; "
        elif motivo == "Falecido":
            paragrafo_unico += "a pessoa procurada já se encontra falecida. "

        if nao_sabe == "Não Conhece ele":
            paragrafo_unico += "não conhece a pessoa procurada, não sabendo informar o local/horário para encontrá-la. "
        elif nao_sabe == "Não sabe informar":
            paragrafo_unico += "que não sabe informar o dia e horário para encontrá-lo(a). "
        elif nao_sabe == "Não sabe seu endereço":
            paragrafo_unico += "que não sabe informar o endereço para encontrá-lo(a). "

        if paradeiro == "Não sabe o paradeiro":
            paragrafo_unico += "não sabe informar seu paradeiro, bem como o local para encontrá-lo. "
        elif paradeiro == "Incerto e Não Sabido":
            paragrafo_unico += "Certifico assim, que, com relação ao presente mandado, endereço fornecido e informações obtidas no local, A PESSOA PROCURADA SE ENCONTRA EM LOCAL INCERTO E NÃO SABIDO. "

    obs_extra = ""
    if condicao == "Chuva":
        obs_extra = "Certifico que a execução da diligência restou dificultada em virtude das adversas condições meteorológicas no momento do ato, caracterizadas por intensa precipitação pluviométrica. Ressalto que tal circunstância, além de elevar significativamente o ruído ambiental comprometendo a audibilidade do chamamento realizado no portão, bem como ocasiona o natural recolhimento dos moradores no interior da residência com janelas e portas cerradas, o que obstaculizou a percepção da minha presença e, consequentemente, impediu o efetivo atendimento. "
    elif condicao == "Local Perigoso":
        obs_extra = "Informo também que o local é conhecidamente de grande periculosidade, o que quase sempre inviabiliza a obtenção de informações, pois os moradores ficam receosos de envolvimento com o processo e suas consequências, onde conversei com alguns vizinhos, que não quiseram se identificar, e ninguém soube informar detalhes sobre o possível horário/local para encontrar a pessoa procurada. "
    elif condicao == "Zona Rural":
        obs_extra = "Informo que o local é uma zona rural com difícil acesso, localização difícil, numeração irregular com muitas casas sem números na porta, o que causa desconforto nos moradores em fornecer informações precisas sobre o local/horário para encontrar a pessoa procurada. "
    elif condicao == "Blocos":
        obs_extra = "Informo também que o local é um condomínio de edifícios com vários blocos de apartamentos em seu interior; possui portaria na entrada do condomínio, mas não existe nenhum porteiro no local em nenhum horário; possui um interfone na entrada que é o único meio de contato com os apartamentos dentro do condomínio, mas aparentemente esse interfone não está funcionando, pois toquei várias vezes e ninguém atendeu; procurei informações com moradores que estavam saindo do condomínio sobre o possível contato com a pessoa procurada, mas ninguém soube informar se o mesmo reside no condomínio dizendo “são muitos moradores e não conhecemos todo mundo”, afirmando não saber informar também o possível horário para encontrá-la. "
    elif condicao == "Medo Processo":
        obs_extra = "Procurei informações com vizinhos sobre o horário/local para encontrar a pessoa procurada, mas os moradores ficam receosos de envolvimento com o processo e suas consequências, onde conversei com alguns vizinhos, que não quiseram se identificar, e ninguém soube informar detalhes sobre o possível horário/local para encontrar a pessoa procurada. "

    if obs_extra or observacoes:
        paragrafo_unico += obs_extra + (" " + observacoes if observacoes else "")

    # ==========================================
    # CRIAÇÃO DO ARQUIVO WORD
    # ==========================================
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 1. CABEÇALHO (IMAGEM)
    if os.path.exists("cabecalho.png"):
        p_img_cabecalho = doc.add_paragraph()
        p_img_cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img_cab = p_img_cabecalho.add_run()
        # Ajuste o tamanho em centímetros conforme necessário
        run_img_cab.add_picture("cabecalho.png", width=Cm(16))
    elif os.path.exists("cabecalho.jpg"):
        p_img_cabecalho = doc.add_paragraph()
        p_img_cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img_cab = p_img_cabecalho.add_run()
        run_img_cab.add_picture("cabecalho.jpg", width=Cm(16))

    # Processo / Ano e Mandado
    if processo:
        texto_processo = f"Processo: {processo}"
        if ano:
            texto_processo += f".{ano}.8.13.0245"
        doc.add_paragraph(texto_processo)
        
    if mandado:
        doc.add_paragraph(f"Mandado nº: {mandado}")
        
    doc.add_paragraph("")

    # Título
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("CERTIDÃO")
    run_titulo.bold = True
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("")

    # Corpo do texto
    p_corpo = doc.add_paragraph(paragrafo_unico.strip())
    p_corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_corpo.paragraph_format.first_line_indent = Pt(35.4) 
    
    doc.add_paragraph("")

    # Fechamento e Data 
    p_fechamento = doc.add_paragraph("Devolvo o mandado para os devidos fins. O referido é verdade. Dou fé.")
    p_fechamento.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    hoje = datetime.date.today()
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    data_extenso = f"Santa Luzia, {hoje.day} de {meses[hoje.month - 1]} de {hoje.year}."
    
    p_data = doc.add_paragraph(data_extenso)
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    
    # 2. ASSINATURA (IMAGEM)
    if os.path.exists("assinatura.png"):
        p_img_assinatura = doc.add_paragraph()
        p_img_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img_ass = p_img_assinatura.add_run()
        # Ajuste o tamanho da assinatura em centímetros
        run_img_ass.add_picture("assinatura.png", width=Cm(5))
    elif os.path.exists("assinatura.jpg"):
        p_img_assinatura = doc.add_paragraph()
        p_img_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img_ass = p_img_assinatura.add_run()
        run_img_ass.add_picture("assinatura.jpg", width=Cm(5))
    
# Dados do Oficial
    p_assinatura = doc.add_paragraph()
    p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_nome = p_assinatura.add_run("Rafael Giardini de Oliveira\n")
    run_nome.bold = True
    run_nome.font.size = Pt(8)
    
    run_cargo = p_assinatura.add_run("Oficial de Justiça Avaliador\n")
    run_cargo.font.size = Pt(8)
    
    run_pjpi = p_assinatura.add_run("PJPI: 22842-9")
    run_pjpi.font.size = Pt(8)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Nomenclatura dinâmica do arquivo
    data_arquivo = hoje.strftime("%d-%m-%Y")
    if processo:
        nome_arquivo = f"Certidao_Negativa_{processo}_{data_arquivo}.docx"
    else:
        nome_arquivo = f"Certidao_Negativa_{data_arquivo}.docx"

    st.success("✅ Certidão gerada!")
    st.download_button(
        label="📥 Baixar Documento Word",
        data=buffer,
        file_name=nome_arquivo,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )