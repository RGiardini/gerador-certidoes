import streamlit as st
import os
import hashlib
import zipfile
import subprocess
import tempfile
from io import BytesIO
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from supabase import create_client, Client



# ==========================================
# 1. FUNÇÃO DE CONVERSÃO PARA PDF
# ==========================================
def converter_docx_para_pdf(docx_bytes):
    """
    Salva o DOCX temporariamente, aciona o LibreOffice oculto para converter 
    e devolve os bytes do PDF gerado.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        caminho_docx = os.path.join(temp_dir, "temp_certidao.docx")
        
        with open(caminho_docx, "wb") as f:
            f.write(docx_bytes)
        
        comando = [
            "libreoffice", "--headless", "--convert-to", "pdf",
            "--outdir", temp_dir, caminho_docx
        ]
        
        subprocess.run(comando, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        caminho_pdf = os.path.join(temp_dir, "temp_certidao.pdf")
        
        if os.path.exists(caminho_pdf):
            with open(caminho_pdf, "rb") as f:
                return f.read()
        return None

# ==========================================
# 2. CONFIGURAÇÃO DA PÁGINA E BANCO DE DADOS
# ==========================================
st.set_page_config(page_title="Sistema de Certidões", layout="wide")

st.markdown("""
    <style>
    /* Oculta elementos padrão do Streamlit */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* Configuração geral do container e espaçamentos modernos */
    .block-container { 
        padding-top: 2rem; 
        padding-bottom: 2rem; 
        max-width: 1200px;
    }
    div[data-testid="stVerticalBlock"] { 
        gap: 1rem !important; 
    }

    /* 🚀 Centralização e refinamento de todos os títulos principais (h1) */
    .main h1, h1[data-testid="stHeader"] { 
        font-size: 26px !important; 
        font-weight: 700;
        text-align: center !important; 
        color: #1E293B;
        margin-top: 0rem !important; 
        margin-bottom: 1.5rem !important; 
        display: block;
        width: 100%;
    }

    /* Criação do efeito de Cartões Flutuantes nos blocos principais */
    div[data-testid="stExpander"], div.stTextInput, div.stSelectbox, div.stRadio {
        background-color: #FFFFFF;
        border-radius: 10px;
        padding: 0.2rem;
    }

    /* Ajustes finos nos inputs de texto */
    input[type="text"], input[type="password"] {
        border-radius: 8px !important;
        border: 1px solid #CBD5E1 !important;
    }
    input[type="text"]:focus, input[type="password"]:focus {
        border-color: #0F172A !important;
        box-shadow: 0 0 0 1px #0F172A !important;
    }

    /* Estilização moderna para os botões primários (Call to Action) */
    button[kind="primary"] {
        background-color: #0F172A !important;
        color: #FFFFFF !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        border: none !important;
        transition: all 0.3s ease !important;
    }
    button[kind="primary"]:hover {
        background-color: #1E293B !important;
        box-shadow: 0 4px 12px rgba(15, 23, 42, 0.2) !important;
    }

    /* Sidebar mais elegante e organizada */
    section[data-testid="stSidebar"] { 
        background-color: #F8FAFC !important;
        border-right: 1px solid #E2E8F0;
        width: 18rem !important; 
    }
    section[data-testid="stSidebar"] div[data-testid="stVerticalBlock"] { 
        gap: 0.8rem !important; 
    }
    </style>
""", unsafe_allow_html=True)

@st.cache_resource
def iniciar_conexao():
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    return create_client(url, key)

supabase: Client = iniciar_conexao()

def gerar_hash_senha(senha):
    return hashlib.sha256(senha.encode()).hexdigest()

# ==========================================
# 3. CONTROLE DE SESSÃO E LOGIN (VIA URL)
# ==========================================

# Se o session_state estiver vazio, tenta resgatar o usuário direto da URL (F5)
if "usuario_logado" not in st.session_state or st.session_state["usuario_logado"] is None:
    usuario_url = st.query_params.get("user")
    if usuario_url:
        st.session_state["usuario_logado"] = usuario_url
    else:
        st.session_state["usuario_logado"] = None

if st.session_state["usuario_logado"] is None:
    st.title("⚖️ Sistema de Certidões - TJMG")
    
    aba_login, aba_cadastro = st.tabs(["Entrar", "Criar Nova Conta"])
    
    with aba_login:
        usuario_login = st.text_input("Usuário:", key="log_usr_input").lower().strip()
        senha_login = st.text_input("Senha:", type="password", key="log_pwd_input")
        
        if st.button("Entrar", type="primary", use_container_width=True, key="btn_entrar"):
            if usuario_login and senha_login:
                resposta = supabase.table("banco_usuarios").select("*").eq("usuario", usuario_login).execute()
                
                if len(resposta.data) > 0:
                    dados_bd = resposta.data[0]
                    senha_criptografada = gerar_hash_senha(senha_login)
                    if dados_bd["senha"] == senha_criptografada:
                        st.session_state["usuario_logado"] = usuario_login
                        # 🚀 Salva o usuário na URL do navegador de forma permanente até o logout
                        st.query_params["user"] = usuario_login
                        st.rerun()
                    else:
                        st.error("Senha incorreta!")
                else:
                    st.error("Usuário não encontrado.")
            else:
                st.warning("Preencha usuário e senha.")
                
    with aba_cadastro:
        pass
        
    st.stop()

# ==========================================
# 4. DADOS DO USUÁRIO E MENU LATERAL
# ==========================================
usuario_atual = st.session_state["usuario_logado"]
resposta_usuario = supabase.table("banco_usuarios").select("*").eq("usuario", usuario_atual).execute()
dados_usuario = resposta_usuario.data[0]

with st.sidebar:
    st.write(f"👤 Olá, **{usuario_atual.title()}**!")
    st.divider()
    
    opcoes_menu = ["📝 Gerar Certidão", "📂 Minhas Certidões", "⚙️ Meu Perfil"]
    if usuario_atual == "10228429":
        opcoes_menu.append("🛡️ Painel do Administrador")
        
    menu = st.radio("Navegação:", opcoes_menu)
    st.divider()
    
    if st.button("Sair (Logout)", key="btn_logout"):
        st.session_state["usuario_logado"] = None
        st.query_params.clear()  # Limpa o usuário da URL
        st.rerun()

# ==========================================
# 5. TELA: MEU PERFIL
# ==========================================
if menu == "⚙️ Meu Perfil":
    st.title("⚙️ Configurar Meu Perfil")
    st.write("Estes dados serão inseridos no final das suas certidões (Fonte tamanho 8).")
    
    novo_nome = st.text_input("Nome Completo:", value=dados_usuario.get("nome", ""), key="input_perfil_nome")
    novo_cargo = st.text_input("Cargo:", value=dados_usuario.get("cargo", ""), key="input_perfil_cargo")
    nova_matricula = st.text_input("Matrícula (ex: PJPI: 12345):", value=dados_usuario.get("matricula", ""), key="input_perfil_matricula")
    
    st.write("**Sua Assinatura (Fundo branco ou transparente):**")
    arquivo_assinatura = st.file_uploader("Envie a foto da sua assinatura", type=["png", "jpg", "jpeg"], key="uploader_perfil")
    
    if st.button("💾 Salvar Perfil", type="primary", use_container_width=True, key="btn_salvar_perfil"):
        supabase.table("banco_usuarios").update({
            "nome": novo_nome,
            "cargo": novo_cargo,
            "matricula": nova_matricula
        }).eq("usuario", usuario_atual).execute()
        
        if arquivo_assinatura is not None:
            try:
                supabase.storage.from_("assinaturas_usuarios").remove([f"{usuario_atual}.png"])
            except:
                pass
            supabase.storage.from_("assinaturas_usuarios").upload(
                file=arquivo_assinatura.getvalue(),
                path=f"{usuario_atual}.png",
                file_options={"content-type": arquivo_assinatura.type}
            )
                
        st.success("✅ Perfil atualizado e salvo na nuvem com sucesso!")
        st.rerun()

# ==========================================
# 6. TELA: MINHAS CERTIDÕES
# ==========================================
elif menu == "📂 Minhas Certidões":
    st.title("📂 Minhas Certidões Salvas")
    st.write("Baixe ou exclua seus arquivos salvos na nuvem.")
    
    try:
        arquivos_nuvem = supabase.storage.from_("certidoes_usuarios").list(usuario_atual)
    except:
        arquivos_nuvem = []
    
    arquivos = [arq for arq in arquivos_nuvem if arq["name"] != ".emptyFolder" and arq["name"] != ""]
    
    if not arquivos:
        st.info("Nenhuma certidão salva ainda.")
    else:
        c_filtro, c_btn1, c_btn2 = st.columns([2, 1.5, 1.5])
        with c_filtro:
            ativar_filtro = st.checkbox("Filtrar por data", key="ativar_filtro_data")
            hoje_real = datetime.datetime.utcnow() - datetime.timedelta(hours=3)
            data_filtro = st.date_input("Escolha a data:", value=hoje_real.date(), format="DD/MM/YYYY", disabled=not ativar_filtro, label_visibility="collapsed")
            
        arquivos_filtrados = []
        for item in arquivos:
            try:
                data_str = item["created_at"].replace("Z", "+00:00")
                data_obj = datetime.datetime.fromisoformat(data_str)
                data_br_obj = data_obj.replace(tzinfo=None) - datetime.timedelta(hours=3)
                data_br_date = data_br_obj.date()
                data_br = data_br_obj.strftime("%d/%m/%Y às %H:%M")
            except:
                data_br_date = None
                data_br = "Data desconhecida"
                
            item['data_br_date'] = data_br_date
            item['data_br'] = data_br
            
            if ativar_filtro and data_br_date != data_filtro:
                continue
            arquivos_filtrados.append(item)
            
        with c_btn1:
            if st.button("✓ Marcar Todos Abaixo", use_container_width=True):
                for arq in arquivos_filtrados:
                    st.session_state[f"chk_file_{arq['name']}"] = True
                st.rerun()
        with c_btn2:
            if st.button("✕ Desmarcar Todos", use_container_width=True):
                for arq in arquivos_filtrados:
                    st.session_state[f"chk_file_{arq['name']}"] = False
                st.rerun()
                
        arquivos_filtrados.sort(key=lambda x: x["created_at"], reverse=True)
        st.divider()
        
        c_sel, c_nome, c_data = st.columns([1, 4, 3])
        c_sel.write("**Selecionar**")
        c_nome.write("**Nome do Arquivo**")
        c_data.write("**Data de Criação**")
        st.divider()
        
        arquivos_selecionados = []
        
        for item in arquivos_filtrados:
            c1, c2, c3 = st.columns([1, 4, 3])
            with c1:
                if st.checkbox("", key=f"chk_file_{item['name']}"):
                    arquivos_selecionados.append(item['name'])
            with c2:
                st.write(item['name'])
            with c3:
                st.write(item['data_br'])
                
        st.divider()
        
        if arquivos_selecionados:
            st.write(f"**{len(arquivos_selecionados)} arquivo(s) selecionado(s)**")
            c_btn1, c_btn2 = st.columns(2)
            
            with c_btn1:
                if st.button("📥 Preparar Download (ZIP)", type="primary", use_container_width=True, key="btn_zip_download"):
                    with st.spinner("Baixando da nuvem..."):
                        zip_buffer = BytesIO()
                        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
                            for arq in arquivos_selecionados:
                                arquivo_bytes = supabase.storage.from_("certidoes_usuarios").download(f"{usuario_atual}/{arq}")
                                zip_file.writestr(arq, arquivo_bytes)
                                
                        st.download_button(
                            label="✔️ Clique aqui para baixar o ZIP",
                            data=zip_buffer.getvalue(),
                            file_name=f"certidoes_{usuario_atual}.zip",
                            mime="application/zip",
                            use_container_width=True,
                            key="download_zip_real"
                        )
            
            with c_btn2:
                if st.button("🗑️ Excluir Selecionadas", use_container_width=True, key="btn_excluir_certidoes"):
                    caminhos_para_excluir = [f"{usuario_atual}/{arq}" for arq in arquivos_selecionados]
                    supabase.storage.from_("certidoes_usuarios").remove(caminhos_para_excluir)
                    st.success("✅ Arquivos excluídos da nuvem com sucesso!")
                    st.rerun()

# ==========================================
# 7. TELA: PAINEL DO ADMINISTRADOR
# ==========================================
elif menu == "🛡️ Painel do Administrador":
    if usuario_atual != "10228429":
        st.error("Acesso restrito apenas ao Administrador.")
        st.stop()
        
    st.title("🛡️ Painel de Administração")
    st.write("Área restrita para gestão de oficiais e auditoria de certidões em nuvem.")
    
    aba_adm1, aba_adm2 = st.tabs(["👥 Gerenciar Usuários", "📊 Auditoria de Certidões Gerais"])
    
    with aba_adm1:
        st.subheader("Oficiais Cadastrados no Sistema")
        res_todos = supabase.table("banco_usuarios").select("usuario, nome, cargo, matricula").execute()
        usuarios_cadastrados = res_todos.data
        
        if usuarios_cadastrados:
            for u in usuarios_cadastrados:
                with st.expander(f"👤 Usuário: {u['usuario']} — Nome: {u.get('nome') or 'Não preenchido'}"):
                    st.write(f"**Cargo:** {u.get('cargo')}")
                    st.write(f"**Matrícula:** {u.get('matricula')}")
                    
                    if u['usuario'] != usuario_atual:
                        if st.button(f"🗑️ Excluir usuário {u['usuario']}", key=f"del_adm_usr_{u['usuario']}"):
                            supabase.table("banco_usuarios").delete().eq("usuario", u['usuario']).execute()
                            st.success(f"Usuário {u['usuario']} removido com sucesso!")
                            st.rerun()
                    else:
                        st.caption("*(Esta é a sua conta de Administrador principal)*")
        else:
            st.info("Nenhum usuário encontrado.")

    with aba_adm2:
        st.subheader("Certidões Geradas por Todos os Oficiais")
        try:
            pastas_usuarios = supabase.storage.from_("certidoes_usuarios").list()
        except:
            pastas_usuarios = []
            
        if not pastas_usuarios:
            st.info("Nenhuma pasta de certidão encontrada na nuvem.")
        else:
            for pasta in pastas_usuarios:
                nome_oficial = pasta["name"]
                if nome_oficial and nome_oficial != ".emptyFolder":
                    st.markdown(f"### 📂 Oficial: `{nome_oficial}`")
                    
                    try:
                        arquivos_oficial = supabase.storage.from_("certidoes_usuarios").list(nome_oficial)
                    except:
                        arquivos_oficial = []
                        
                    certioes_validas = [f for f in arquivos_oficial if f["name"] != ".emptyFolder" and f["name"] != ""]
                    
                    if not certioes_validas:
                        st.caption("Nenhuma certidão gerada por este oficial ainda.")
                    else:
                        for arq in certioes_validas:
                            c_arq_nome, c_btn_dl, c_btn_del = st.columns([4, 2, 2])
                            
                            with c_arq_nome:
                                st.text(arq["name"])
                                
                            with c_btn_dl:
                                if st.button("📥 Baixar", key=f"dl_adm_f_{nome_oficial}_{arq['name']}", use_container_width=True):
                                    file_bytes = supabase.storage.from_("certidoes_usuarios").download(f"{nome_oficial}/{arq['name']}")
                                    
                                    # Ajuste Dinâmico de MIME Type
                                    mime_tipo = "application/pdf" if arq["name"].endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    
                                    st.download_button(
                                        label="Confirmar",
                                        data=file_bytes,
                                        file_name=arq["name"],
                                        mime=mime_tipo,
                                        key=f"btn_dl_real_{nome_oficial}_{arq['name']}"
                                    )
                                    
                            with c_btn_del:
                                if st.button("🗑️ Excluir", key=f"del_adm_f_{nome_oficial}_{arq['name']}", use_container_width=True):
                                    supabase.storage.from_("certidoes_usuarios").remove([f"{nome_oficial}/{arq['name']}"])
                                    st.success("Excluído!")
                                    st.rerun()
                    st.divider()

# ==========================================
# 8. TELA: GERADOR DE CERTIDÃO
# ==========================================
elif menu == "📝 Gerar Certidão":
    st.title("Gerador de Certidões - TJMG")
    
    if not dados_usuario.get("nome"):
        st.warning("⚠️ Você ainda não configurou seu perfil! Vá em 'Meu Perfil' antes de gerar certidões.")
        st.stop()

    # --- ESCOLHAS PRINCIPAIS ---
    c_tipo, c_formato = st.columns([3, 1])
    with c_tipo:
        tipo_certidao = st.selectbox(
            "Selecione o Modelo de Certidão:", 
            [
                "Certidão Negativa Detalhada (Antiga)", 
                "Certidão Negativa Detalhada (Nova Versão)",
                "Certidão Negativa Simples (Opções Rápidas)", 
                "Certidão Positiva",
                "Certidão Positiva por Hora Certa"
            ]
        )
    with c_formato:
        formato_saida = st.radio(
            "Formato de exportação:", 
            ["Word (.docx)", "PDF (.pdf)"], 
            key="formato_global"
        )
        
    st.divider()

    # --- CAMPOS COMPARTILHADOS (CABEÇALHO) ---
    c_mandado, c_proc, c_ano, c_comarca = st.columns([1, 2.5, 1, 1])
    
    with c_mandado:
        mandado = st.text_input("Mandado:", placeholder="Ex: 01", key="mandado_geral")
    with c_proc:
        processo = st.text_input("Processo:", placeholder="Ex: 4400281-16", key="processo_geral")
    with c_ano:
        ano = st.text_input("Ano:", value="2026", placeholder="Ex: 2026", key="ano_geral")
    with c_comarca:
        comarca = st.text_input("Cód. Comarca:", value="0245", placeholder="Ex: 0245", key="comarca_geral")

    c_end, c_pes = st.columns(2)
    with c_end:
        endereco = st.text_input("Endereço (opcional):", placeholder="Se vazio: 'informado no mesmo'", key="endereco_geral")
    with c_pes:
        pessoa = st.text_input("Pessoa procurada:", placeholder="Deixe vazio para termo genérico", key="pessoa_geral")

    st.markdown("---")
    st.subheader("Data do Documento e Diligências")
    
    hoje_real = datetime.datetime.utcnow() - datetime.timedelta(hours=3)
    data_certidao = st.date_input("Data que sairá no rodapé da certidão:", value=hoje_real.date(), format="DD/MM/YYYY", key="data_certidao_geral")
    
    st.write("**Informe os Dias e Horários (o 'h' será adicionado automaticamente se esquecer):**")
    
    cd1, cd2, cd3 = st.columns(3)

    with cd1:
        st.write("**Diligência 1**")
        d1 = st.text_input("Dia 1", placeholder="Ex: 08/08", key="d1_geral")
        h1 = st.text_input("Hora 1", placeholder="Ex: 14:55", key="h1_geral")
        
    with cd2:
        st.write("**Diligência 2**")
        d2 = st.text_input("Dia 2", placeholder="Ex: 11/08", key="d2_geral")
        h2 = st.text_input("Hora 2", placeholder="Ex: 16:58", key="h2_geral")
        
    with cd3:
        st.write("**Diligência 3**")
        d3 = st.text_input("Dia 3", placeholder="Ex: 12/08", key="d3_geral")
        h3 = st.text_input("Hora 3", placeholder="Ex: 11:15", key="h3_geral")

    st.divider()

    # ==========================================
    # OPÇÃO A: CERTIDÃO DETALHADA ANTIGA
    # ==========================================
    if tipo_certidao == "Certidão Negativa Detalhada (Antiga)":
        
        if st.session_state.get('limpar_detalhada'):
            for k in list(st.session_state.keys()):
                if k.startswith(("mot_det_", "rel_det_", "ns_det_", "cert_")) or k in ["nao_loc_dest", "nao_loc_bens"]:
                    st.session_state[k] = False 
                elif k in ["nome_inf_det", "sabe_tel_det", "sabe_end_det", "obs_livres_det"]:
                    st.session_state[k] = ""
            st.session_state['limpar_detalhada'] = False

        st.write("**Deixei de cumprir a ordem descrita uma vez que:**")
        sit_c1, sit_c2 = st.columns(2)
        with sit_c1:
            nao_loc_dest = st.checkbox("O destinatário não foi localizado", key="nao_loc_dest")
        with sit_c2:
            nao_loc_bens = st.checkbox("Bem(ns) não localizado(s)", key="nao_loc_bens")

        motivos_selecionados = []
        with st.expander("📌 Selecionar Motivos Detalhados", expanded=False):
            motivos_list = [
                "mudou-se", "não reside no local", "não foi localizada", "é desconhecido", "dificilmente fica ali", "trabalha em tempo integral",
                "não trabalha no local", "está viajando", "local inabitado", "antigo inquilino", 
                "antigo morador", "antigo proprietário", "rotatividade de inquilinos",
                "Repassado para terceiros", "internado", "transferido", "encontra-se preso",
                "faleceu", "faliu", "não exerce atividades", "local fechado", 
                "número não localizado", "rua/av não localizada", "ap/bloco não localizado", 
                "aparece esporadicamente", "utiliza endereço para correspondências",
                "sem condições psíquicas de entender conteúdo mandado",
                "guarnecem a residência amparados pela Lei 8.009/90",
                "são insuficientes para saldar o débito"
            ]
            cols_mot = st.columns(3)
            for idx, m in enumerate(motivos_list):
                with cols_mot[idx % 3]:
                    if st.checkbox(m, key=f"mot_det_{idx}"):
                        motivos_selecionados.append(m)

        st.markdown("---")
        relacoes_selecionadas = []
        nao_sabe_selecionados = []
        sabe_tel = ""
        sabe_end = ""
        
        with st.expander("👤 Informações sobre o Informante", expanded=False):
            nome_inf_det = st.text_input("Nome do Sr(a):", placeholder="Vazio se não houver informante", key="nome_inf_det")
            st.caption("Relação / Qualidade:")
            relacoes_list = [
                "morador", "proprietário", "inquilino", "funcionário", "vizinho", "pai", "mãe",
                "padrasto", "madrasta", "filho", "irmão", "tio", "avô(ó)", "neto", "sobrinho",
                "primo", "transeunte", "viúvo", "ex", "esposo", "companheiro", "sogro", "enteado",
                "genro", "nora", "cunhado", "concunhado", "amigo"
            ]
            cols_rel = st.columns(4)
            for idx, r in enumerate(relacoes_list):
                with cols_rel[idx % 4]:
                    if st.checkbox(r, key=f"rel_det_{idx}"):
                        relacoes_selecionadas.append(r)

            st.write("**Não sabendo o informante indicar:**")
            nao_sabe_list = [
                "endereço completo", "paradeiro da pessoa procurada", "o dia/horário exato para encontrá-lo(a)", 
                "telefone", "dia/horário de retorno", "o presídio", 
                "dados do óbito", "previsão de alta"
            ]
            cols_ns = st.columns(3)
            for idx, ns in enumerate(nao_sabe_list):
                with cols_ns[idx % 3]:
                    if st.checkbox(ns, key=f"ns_det_{idx}"):
                        nao_sabe_selecionados.append(ns)

            st.write("**Sabendo o informante indicar:**")
            sabe_tel = st.text_input("Telefone indicado:", key="sabe_tel_det")
            sabe_end = st.text_input("Endereço correto indicado:", key="sabe_end_det")
        
        with st.expander("📝 Certificações Adicionais", expanded=False):
            cert_extras = []
            c_extra1, c_extra2 = st.columns(2)
            with c_extra1:
                if st.checkbox("Procurei informações com moradores", key="cert_vizinhos_det"):
                    cert_extras.append("procurei obter informações junto aos moradores/vizinhos locais e não obtive êxito")
                if st.checkbox("Cópia do mandado com informante", key="cert_copia_det"):
                    cert_extras.append("devido à importância do mandado e da dificuldade de encontrar a pessoa procurada, deixei a cópia do mandado com o(a) senhor(a) acima mencionado(a)")
                if st.checkbox("Local Perigoso", key="cert_perigoso_det"):
                    cert_extras.append("o local é conhecidamente de grande periculosidade, o que quase sempre inviabiliza a obtenção de informações, pois os moradores ficam receosos de envolvimento com o processo e suas consequências, onde conversei com alguns vizinhos, que não quiseram se identificar, e ninguém soube informar detalhes sobre o possível horário/local para encontrar a pessoa procurada")
                if st.checkbox("Medo do Processo", key="cert_medo_det"):
                    cert_extras.append("os moradores ficam receosos de envolvimento com o processo e suas consequências, onde conversei com alguns vizinhos, que não quiseram se identificar, e ninguém soube informar detalhes sobre o possível horário/local para encontrar a pessoa procurada")
            with c_extra2:
                if st.checkbox("Apenas bens domésticos", key="cert_moveis_det"):
                    cert_extras.append("o imóvel é residencial e contém apenas móveis e utensílios domésticos")
                if st.checkbox("Zona Rural", key="cert_rural_det"):
                    cert_extras.append(" local é uma zona rural com difícil acesso, localização difícil, numeração irregular com muitas casas sem números na porta, o que causa desconforto nos moradores em fornecer informações precisas sobre o local/horário para encontrar a pessoa procurada")
                if st.checkbox("Condomínio de Blocos", key="cert_blocos_det"):
                    cert_extras.append("o local é um condomínio de edifícios com vários blocos de apartamentos em seu interior; possui portaria na entrada do condomínio, mas não existe nenhum porteiro no local em nenhum horário; possui um interfone na entrada que é o único meio de contato com os apartamentos dentro do condomínio, mas aparentemente esse interfone não está funcionando, pois toquei várias vezes e ninguém atendeu; procurei informações com moradores que estavam saindo do condomínio sobre o possível contato com a pessoa procurada, mas ninguém soube informar se o mesmo reside no condomínio dizendo “são muitos moradores e não conhecemos todo mundo”, afirmando não saber informar também o possível horário para encontrá-la")
                if st.checkbox("Chuva Forte", key="cert_chuva_det"):
                    cert_extras.append("a execução da diligência restou dificultada em virtude das adversas condições meteorológicas no momento do ato, caracterizadas por intensa precipitação pluviométrica. Ressalto que tal circunstância, além de elevar significativamente o ruído ambiental comprometendo a audibilidade do chamamento realizado no portão, bem como ocasiona o natural recolhimento dos moradores no interior da residência com janelas e portas cerradas, o que obstaculizou a percepção da minha presença e, consequentemente, impediu o efetivo atendimento")

            observacoes_det = st.text_area("Observações Livres:", key="obs_livres_det")

        st.divider()

        if st.button("Salvar na Nuvem / Gerar Documento", type="primary", use_container_width=True, key="btn_gerar_docx_det"):
            with st.spinner("Gerando detalhada..."):
                dias_validos = [d for d in [d1, d2, d3] if d]
                horas_cruas = [h for h in [h1, h2, h3] if h]
                horas_validas = []
                for h in horas_cruas:
                    h_limpo = h.strip()
                    if h_limpo and not h_limpo.lower().endswith(('h', 'hs')):
                        h_limpo += 'hs'
                    horas_validas.append(h_limpo)

                texto_data_hora = ""
                if len(dias_validos) == 1:
                    texto_data_hora = f", por volta das {horas_validas[0]}, do dia {dias_validos[0]},"
                elif len(dias_validos) > 1:
                    str_horas = ", ".join(horas_validas[:-1]) + f" e {horas_validas[-1]}"
                    str_dias = ", ".join(dias_validos[:-1]) + f" e {dias_validos[-1]}"
                    texto_data_hora = f", por volta das {str_horas}, dos dias {str_dias}, respectivamente,"
                
                txt_endereco = f"à {endereco}" if endereco else "ao endereço/local/região/bairro indicado(a)"
                txt_pessoa = f" em face de {pessoa}" if pessoa else ""
                paragrafo = f"Certifico que, em cumprimento ao mandado anexo, desloquei-me {txt_endereco}{texto_data_hora} onde deixei de cumprir a ordem descrita{txt_pessoa}, uma vez que "
                
                sits = []
                if nao_loc_dest: sits.append("o destinatário do mandado não foi localizado")
                if nao_loc_bens: sits.append("o(s) bem(ns) indicados não foi(ram) localizado(s)")
                paragrafo += " e ".join(sits) + ". " if sits else "não foi possível a sua realização. "
                
                if motivos_selecionados:
                    frases_motivos = []
                    for m in motivos_selecionados:
                        if m == "mudou-se": frases_motivos.append("a pessoa procurada não reside mais no local, tendo se mudado")
                        elif m == "não reside no local": frases_motivos.append("a pessoa procurada não reside no local indicado")
                        elif m == "não foi localizada": frases_motivos.append("a pessoa procurada não foi localizada")
                        elif m == "faliu": frases_motivos.append("a empresa procurada faliu ou encerrou suas atividades")
                        else: frases_motivos.append(f"a pessoa procurada {m}")
                    
                    if len(frases_motivos) > 1:
                        texto_motivos = ", e que ".join([", ".join(frases_motivos[:-1]), frases_motivos[-1]])
                    else:
                        texto_motivos = frases_motivos[0]
                    paragrafo += f"Constatou-se na diligência que {texto_motivos}. "
                
                if nome_inf_det or relacoes_selecionadas:
                    txt_informante = f"pelo(a) Sr(a). {nome_inf_det}" if nome_inf_det else "por pessoa não identificada"
                    rel_str = f", na qualidade de {', '.join(relacoes_selecionadas)}," if relacoes_selecionadas else ""
                    paragrafo += f"Conforme informações prestadas no local {txt_informante}{rel_str} "
                    
                    if nao_sabe_selecionados:
                        if len(nao_sabe_selecionados) > 1:
                            texto_ns = ", ".join(nao_sabe_selecionados[:-1]) + f" e nem {nao_sabe_selecionados[-1]}"
                        else:
                            texto_ns = nao_sabe_selecionados[0]
                        paragrafo += f"este(a) declarou não saber informar {texto_ns}. "
                    else:
                        paragrafo += "este(a) prestou as devidas informações no local. "
                        
                    if sabe_tel or sabe_end:
                        sabes_list = []
                        if sabe_tel: sabes_list.append(f"o telefone de contato {sabe_tel}")
                        if sabe_end: sabes_list.append(f"o endereço atual/correto sendo {sabe_end}")
                        paragrafo += f"Por outro lado, o informante soube indicar {' e '.join(sabes_list)}. "
                        
                if cert_extras: paragrafo += f"Certifico também que {'; '.join(cert_extras)}. "
                if observacoes_det: paragrafo += f"{observacoes_det.strip()} "
                
                doc = Document(); style = doc.styles['Normal']; font = style.font; font.name = 'Times New Roman'; font.size = Pt(12)
                try:
                    cabecalho_bytes = supabase.storage.from_("imagens_sistema").download("cabecalho.png")
                    p_img_cabecalho = doc.add_paragraph(); p_img_cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_cabecalho.add_run().add_picture(BytesIO(cabecalho_bytes), width=Cm(16))
                except Exception as e:
                    st.error(f"Erro do Supabase: {e}")
                if processo:
                    texto_processo = f"Processo: {processo}"
                    if ano: texto_processo += f".{ano}.8.13.{comarca}"
                    doc.add_paragraph(texto_processo)
                if mandado: doc.add_paragraph(f"Mandado nº: {mandado}")
                doc.add_paragraph(""); p_titulo = doc.add_paragraph(); run_titulo = p_titulo.add_run("CERTIDÃO"); run_titulo.bold = True; run_titulo.font.size = Pt(16); p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("")
                doc.add_paragraph(paragrafo.strip()).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; doc.paragraphs[-1].paragraph_format.first_line_indent = Pt(35.4); doc.add_paragraph("")
                doc.add_paragraph("Devolvo o mandado para os devidos fins. É verdade. Dou fé.").alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
                doc.add_paragraph(f"Santa Luzia, {data_certidao.day} de {meses[data_certidao.month - 1]} de {data_certidao.year}.").alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                
                try:
                    assinatura_bytes = supabase.storage.from_("assinaturas_usuarios").download(f"{usuario_atual}.png")
                    p_img_assinatura = doc.add_paragraph(); p_img_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img_assinatura.add_run().add_picture(BytesIO(assinatura_bytes), width=Cm(5))
                except: pass 
                p_assinatura = doc.add_paragraph(); p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_nome = p_assinatura.add_run(f"{dados_usuario['nome']}\n"); run_nome.bold = True; run_nome.font.size = Pt(8)
                run_cargo = p_assinatura.add_run(f"{dados_usuario['cargo']}\n"); run_cargo.font.size = Pt(8)
                run_matricula = p_assinatura.add_run(f"{dados_usuario['matricula']}"); run_matricula.font.size = Pt(8)
                
                buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
                docx_bytes = buffer.getvalue()
                
                data_arquivo = hoje_real.strftime("%d-%m-%Y_%Hh%M")
                nome_base = f"Certidao_Negativa_Antiga_{processo}_{data_arquivo}"
                
                if formato_saida == "PDF (.pdf)":
                    arquivo_final_bytes = converter_docx_para_pdf(docx_bytes)
                    nome_final = nome_base + ".pdf"
                    mime_final = "application/pdf"
                    if not arquivo_final_bytes:
                        st.error("Erro na conversão PDF. Baixando DOCX.")
                        arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                supabase.storage.from_("certidoes_usuarios").upload(file=arquivo_final_bytes, path=f"{usuario_atual}/{nome_final}", file_options={"content-type": mime_final})
                
                st.session_state['doc_pronto_bytes_a1'] = arquivo_final_bytes
                st.session_state['doc_pronto_nome_a1'] = nome_final
                st.session_state['doc_pronto_mime_a1'] = mime_final
                st.session_state['piscar_tela'] = True
                st.session_state['limpar_detalhada'] = True
                st.rerun()

        if 'doc_pronto_bytes_a1' in st.session_state:
            if st.session_state.get('piscar_tela'):
                st.balloons(); st.toast("✅ Certidão gerada!", icon="🎉")
                st.session_state['piscar_tela'] = False
            st.success("✅ Certidão salva na Nuvem!")
            st.download_button("📥 Baixar Arquivo", data=st.session_state['doc_pronto_bytes_a1'], file_name=st.session_state['doc_pronto_nome_a1'], mime=st.session_state['doc_pronto_mime_a1'], type="primary", use_container_width=True)

    # ==========================================
    # OPÇÃO A.2: CERTIDÃO DETALHADA (NOVA VERSÃO)
    # ==========================================
    elif tipo_certidao == "Certidão Negativa Detalhada (Nova Versão)":
        
        st.info("💡 Versão com nova lógica de construção de texto para maior coesão gramatical.")

        if st.session_state.get('limpar_detalhada_nova'):
            for k in list(st.session_state.keys()):
                if k.startswith(("mot_detn_", "rel_detn_", "ns_detn_", "cert_n_")) or k in ["nao_loc_dest_n", "nao_loc_bens_n"]:
                    st.session_state[k] = False 
                elif k in ["nome_inf_det_n", "sabe_tel_det_n", "sabe_end_det_n", "obs_livres_det_n"]:
                    st.session_state[k] = ""
            st.session_state['limpar_detalhada_nova'] = False

        st.write("**Deixei de cumprir a ordem descrita uma vez que:**")
        sit_c1, sit_c2 = st.columns(2)
        with sit_c1:
            nao_loc_dest_n = st.checkbox("O destinatário não foi localizado", key="nao_loc_dest_n")
        with sit_c2:
            nao_loc_bens_n = st.checkbox("Bem(ns) não localizado(s)", key="nao_loc_bens_n")

        motivos_selecionados = []
        with st.expander("📌 Selecionar Motivos Detalhados", expanded=False):
            motivos_list = [
                "mudou-se", "não reside no local", "não foi localizada", "é desconhecido", "dificilmente fica ali", "trabalha em tempo integral",
                "não trabalha no local", "está viajando", "local inabitado", "antigo inquilino", 
                "antigo morador", "antigo proprietário", "rotatividade de inquilinos",
                "Repassado para terceiros", "internado", "transferido", "encontra-se preso",
                "faleceu", "faliu", "não exerce atividades", "local fechado", 
                "número não localizado", "rua/av não localizada", "ap/bloco não localizado", 
                "aparece esporadicamente", "utiliza endereço para correspondências",
                "sem condições psíquicas de entender conteúdo mandado",
                "guarnecem a residência amparados pela Lei 8.009/90",
                "são insuficientes para saldar o débito"
            ]
            cols_mot = st.columns(3)
            for idx, m in enumerate(motivos_list):
                with cols_mot[idx % 3]:
                    if st.checkbox(m, key=f"mot_detn_{idx}"):
                        motivos_selecionados.append(m)

        st.markdown("---")
        relacoes_selecionadas = []
        nao_sabe_selecionados = []
        sabe_tel = ""
        sabe_end = ""
        
        with st.expander("👤 Informações sobre o Informante", expanded=False):
            nome_inf_det_n = st.text_input("Nome do Sr(a):", placeholder="Vazio se não houver informante", key="nome_inf_det_n")
            st.caption("Relação / Qualidade:")
            relacoes_list = [
                "morador", "proprietário", "inquilino", "funcionário", "vizinho", "pai", "mãe",
                "padrasto", "madrasta", "filho", "irmão", "tio", "avô(ó)", "neto", "sobrinho",
                "primo", "transeunte", "viúvo", "ex", "esposo", "companheiro", "sogro", "enteado",
                "genro", "nora", "cunhado", "concunhado", "amigo"
            ]
            cols_rel = st.columns(4)
            for idx, r in enumerate(relacoes_list):
                with cols_rel[idx % 4]:
                    if st.checkbox(r, key=f"rel_detn_{idx}"):
                        relacoes_selecionadas.append(r)

            st.write("**Não sabendo o informante indicar:**")
            nao_sabe_list = [
                "endereço completo", "paradeiro da pessoa procurada", "o dia/horário exato para encontrá-lo(a)", 
                "telefone", "dia/horário de retorno", "o presídio", 
                "dados do óbito", "previsão de alta"
            ]
            cols_ns = st.columns(3)
            for idx, ns in enumerate(nao_sabe_list):
                with cols_ns[idx % 3]:
                    if st.checkbox(ns, key=f"ns_detn_{idx}"):
                        nao_sabe_selecionados.append(ns)

            st.write("**Sabendo o informante indicar:**")
            sabe_tel = st.text_input("Telefone indicado:", key="sabe_tel_det_n")
            sabe_end = st.text_input("Endereço correto indicado:", key="sabe_end_det_n")
        
        with st.expander("📝 Certificações Adicionais", expanded=False):
            cert_extras = []
            c_extra1, c_extra2 = st.columns(2)
            with c_extra1:
                if st.checkbox("Procurei inf. moradores", key="cert_n_vizinhos_det"): cert_extras.append("busquei informações adicionais com moradores/vizinhos, não obtendo êxito")
                if st.checkbox("Cópia mandado informante", key="cert_n_copia_det"): cert_extras.append("deixei a cópia do mandado com o(a) informante")
                if st.checkbox("Local Perigoso", key="cert_n_perigoso_det"): cert_extras.append("trata-se de local com reconhecida periculosidade")
                if st.checkbox("Medo do Processo", key="cert_n_medo_det"): cert_extras.append("os moradores locais demonstraram claro receio de envolvimento")
            with c_extra2:
                if st.checkbox("Apenas bens domésticos", key="cert_n_moveis_det"): cert_extras.append("o imóvel guarnece apenas bens domésticos")
                if st.checkbox("Zona Rural", key="cert_n_rural_det"): cert_extras.append("o local trata-se de zona rural de difícil acesso")
                if st.checkbox("Cond. Blocos s/ Porteiro", key="cert_n_blocos_det"): cert_extras.append("trata-se de condomínio de blocos com portaria desguarnecida")
                if st.checkbox("Chuva Forte", key="cert_n_chuva_det"): cert_extras.append("diligência prejudicada pelas adversas condições meteorológicas")

            observacoes_det = st.text_area("Observações Livres:", key="obs_livres_det_n")

        st.divider()

        if st.button("Salvar na Nuvem / Gerar Documento", type="primary", use_container_width=True, key="btn_gerar_docx_det_n"):
            with st.spinner("Construindo certidão e preparando arquivo..."):
                
                dias_validos = [d for d in [d1, d2, d3] if d]
                horas_cruas = [h for h in [h1, h2, h3] if h]
                horas_validas = []
                for h in horas_cruas:
                    h_limpo = h.strip()
                    if h_limpo and not h_limpo.lower().endswith(('h', 'hs')):
                        h_limpo += 'h' 
                    horas_validas.append(h_limpo)

                texto_data_hora = ""
                if len(dias_validos) == 1: texto_data_hora = f"no dia {dias_validos[0]}, por volta das {horas_validas[0]},"
                elif len(dias_validos) > 1:
                    str_horas = ", ".join(horas_validas[:-1]) + f" e {horas_validas[-1]}"
                    str_dias = ", ".join(dias_validos[:-1]) + f" e {dias_validos[-1]}"
                    texto_data_hora = f"nos dias {str_dias}, por volta das {str_horas}, respectivamente,"

                txt_endereco = f"à {endereco}" if endereco else "ao endereço indicado no mandado"
                txt_pessoa = f" em face de {pessoa}" if pessoa else ""
                
                paragrafo = f"Certifico e dou fé que, em cumprimento ao mandado anexo, dirigi-me {txt_endereco}, {texto_data_hora} ocasião em que deixei de cumprir a ordem descrita{txt_pessoa}, uma vez que "
                
                sits = []
                if nao_loc_dest_n: sits.append("o destinatário não foi localizado")
                if nao_loc_bens_n: sits.append("o(s) bem(ns) indicados não foi(ram) localizado(s)")
                paragrafo += " e ".join(sits) + ". " if sits else "não foi possível a sua realização. "

                if motivos_selecionados:
                    frases_motivos = []
                    for m in motivos_selecionados:
                        if m == "mudou-se": frases_motivos.append("a pessoa procurada não reside mais no local")
                        elif m == "local fechado": frases_motivos.append("o imóvel encontrava-se fechado nas ocasiões das diligências")
                        elif m == "faliu": frases_motivos.append("a empresa faliu ou encerrou suas atividades")
                        else: frases_motivos.append(f"a pessoa procurada {m}")
                    
                    if len(frases_motivos) > 1: texto_motivos = ", e que ".join([", ".join(frases_motivos[:-1]), frases_motivos[-1]])
                    else: texto_motivos = frases_motivos[0]
                    paragrafo += f"Constatou-se na diligência que {texto_motivos}. "

                if nome_inf_det_n or relacoes_selecionadas or nao_sabe_selecionados or sabe_tel or sabe_end:
                    if nome_inf_det_n:
                        txt_informante = f"pelo(a) Sr(a). {nome_inf_det_n}"
                        txt_informante += f", na qualidade de {', '.join(relacoes_selecionadas)}," if relacoes_selecionadas else ","
                    else:
                        if relacoes_selecionadas: txt_informante = f"por um(a) {', '.join(relacoes_selecionadas)} que preferiu não se identificar,"
                        else: txt_informante = "por pessoa no local que preferiu não se identificar,"
                            
                    paragrafo += f"Conforme informações prestadas no local {txt_informante} "
                    
                    if nao_sabe_selecionados:
                        if len(nao_sabe_selecionados) > 1: texto_ns = ", ".join(nao_sabe_selecionados[:-1]) + f" e nem {nao_sabe_selecionados[-1]}"
                        else: texto_ns = nao_sabe_selecionados[0]
                        paragrafo += f"este(a) declarou não saber informar {texto_ns}. "
                    else:
                        paragrafo += "nada mais sendo declarado. "
                        
                    if sabe_tel or sabe_end:
                        sabes_list = []
                        if sabe_tel: sabes_list.append(f"o telefone de contato {sabe_tel}")
                        if sabe_end: sabes_list.append(f"o endereço atual como sendo: {sabe_end}")
                        paragrafo += f"Por outro lado, a referida pessoa soube indicar {' e '.join(sabes_list)}. "

                if cert_extras: paragrafo += f"Certifico ainda que {'; '.join(cert_extras)}. "
                if observacoes_det: paragrafo += f"{observacoes_det.strip()} "

                doc = Document(); style = doc.styles['Normal']; font = style.font; font.name = 'Times New Roman'; font.size = Pt(12)
                try:
                    cabecalho_bytes = supabase.storage.from_("imagens_sistema").download("cabecalho.png")
                    p_img_cabecalho = doc.add_paragraph(); p_img_cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_cabecalho.add_run().add_picture(BytesIO(cabecalho_bytes), width=Cm(16))
                except: pass
                if processo:
                    texto_processo = f"Processo: {processo}"
                    if ano: texto_processo += f".{ano}.8.13.{comarca}"
                    doc.add_paragraph(texto_processo)
                if mandado: doc.add_paragraph(f"Mandado nº: {mandado}")
                doc.add_paragraph(""); p_titulo = doc.add_paragraph(); run_titulo = p_titulo.add_run("CERTIDÃO"); run_titulo.bold = True; run_titulo.font.size = Pt(16); p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                doc.add_paragraph(paragrafo.strip()).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; doc.paragraphs[-1].paragraph_format.first_line_indent = Pt(35.4); doc.add_paragraph("")
                doc.add_paragraph("Devolvo o mandado para os devidos fins. É verdade. Dou fé.").alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
                doc.add_paragraph(f"Santa Luzia, {data_certidao.day} de {meses[data_certidao.month - 1]} de {data_certidao.year}.").alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                
                try:
                    assinatura_bytes = supabase.storage.from_("assinaturas_usuarios").download(f"{usuario_atual}.png")
                    p_img_assinatura = doc.add_paragraph(); p_img_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_assinatura.add_run().add_picture(BytesIO(assinatura_bytes), width=Cm(5))
                except: pass 
                p_assinatura = doc.add_paragraph(); p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_nome = p_assinatura.add_run(f"{dados_usuario['nome']}\n"); run_nome.bold = True; run_nome.font.size = Pt(8)
                run_cargo = p_assinatura.add_run(f"{dados_usuario['cargo']}\n"); run_cargo.font.size = Pt(8)
                run_matricula = p_assinatura.add_run(f"{dados_usuario['matricula']}"); run_matricula.font.size = Pt(8)
                
                buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
                docx_bytes = buffer.getvalue()
                
                data_arquivo = hoje_real.strftime("%d-%m-%Y_%Hh%M")
                nome_base = f"Certidao_Nova_Detalhada_{processo}_{data_arquivo}"
                
                if formato_saida == "PDF (.pdf)":
                    arquivo_final_bytes = converter_docx_para_pdf(docx_bytes)
                    nome_final = nome_base + ".pdf"
                    mime_final = "application/pdf"
                    if not arquivo_final_bytes:
                        st.error("Erro na conversão PDF. Baixando DOCX.")
                        arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                supabase.storage.from_("certidoes_usuarios").upload(file=arquivo_final_bytes, path=f"{usuario_atual}/{nome_final}", file_options={"content-type": mime_final})
                
                st.session_state['doc_pronto_bytes_a2'] = arquivo_final_bytes
                st.session_state['doc_pronto_nome_a2'] = nome_final
                st.session_state['doc_pronto_mime_a2'] = mime_final
                st.session_state['piscar_tela'] = True
                st.session_state['limpar_detalhada_nova'] = True
                st.rerun()

        if 'doc_pronto_bytes_a2' in st.session_state:
            if st.session_state.get('piscar_tela'):
                st.balloons(); st.toast("✅ Nova versão gerada!", icon="🎉")
                st.session_state['piscar_tela'] = False
            st.success("✅ Certidão salva na Nuvem!")
            st.download_button("📥 Baixar Arquivo", data=st.session_state['doc_pronto_bytes_a2'], file_name=st.session_state['doc_pronto_nome_a2'], mime=st.session_state['doc_pronto_mime_a2'], type="primary", use_container_width=True)
            
    # ==========================================
    # OPÇÃO B: CERTIDÃO SIMPLES
    # ==========================================
    elif tipo_certidao == "Certidão Negativa Simples (Opções Rápidas)":
        
        if st.session_state.get('limpar_simples'):
            for k in ["sit_radio_simples", "obteve_inf_radio_simples", "motivo_radio_simples", "naosabe_radio_simples", "paradeiro_radio_simples", "condicao_radio_simples"]:
                st.session_state[k] = None
            for k in ["nome_inf_input_simples", "obs_simples"]:
                st.session_state[k] = ""
            st.session_state['limpar_simples'] = False
            
        st.subheader("Situação Principal")
        situacao_simples = st.radio("Selecione:", ["Local Fechado", "Pessoa Não Encontrada", "Não Localizei a Pessoa"], index=None, horizontal=True, key="sit_radio_simples")

        st.markdown("---")
        obteve_inf_simples = st.radio("Obteve Informações?", ["Sim", "Não", "NQI"], index=None, horizontal=True, key="obteve_inf_radio_simples")
        nome_inf_simples = st.text_input("Nome do informante:", disabled=(obteve_inf_simples != "Sim"), key="nome_inf_input_simples")

        st.markdown("---")
        st.write("**Detalhes das Informações Obtidas:**")
        c1, c2, c3 = st.columns(3)
        with c1: motivo_simples = st.radio("Motivo:", ["Mudou-se", "Não Reside no Local", "Não fica ali", "Não trabalha ali", "Falecido"], index=None, key="motivo_radio_simples")
        with c2: nao_sabe_simples = st.radio("O que não sabe?", ["Não Conhece", "Não sabe informar", "Não sabe endereço"], index=None, key="naosabe_radio_simples")
        with c3: paradeiro_simples = st.radio("Paradeiro:", ["Não sabe o paradeiro", "Incerto e Não Sabido"], index=None, key="paradeiro_radio_simples")

        st.markdown("---")
        st.write("**Condições Extras**")
        condicao_simples = st.radio("Selecione:", ["Local Perigoso", "Medo Processo", "Zona Rural", "Blocos s/ Porteiro", "Chuva"], index=None, horizontal=True, key="condicao_radio_simples")

        st.markdown("---")
        observacoes_simples = st.text_area("Observações Extras:", height=60, key="obs_simples")
        st.divider()

        if st.button("Salvar na Nuvem / Gerar Documento", type="primary", use_container_width=True, key="btn_gerar_simples"):
            with st.spinner("Construindo certidão simples e salvando na nuvem..."):
                dias_validos = [d for d in [d1, d2, d3] if d]
                horas_cruas = [h for h in [h1, h2, h3] if h]
                horas_validas = []
                for h in horas_cruas:
                    h_limpo = h.strip()
                    if h_limpo and not h_limpo.lower().endswith(('h', 'hs')): h_limpo += 'hs'
                    horas_validas.append(h_limpo)
                
                texto_data_hora = ""
                if len(dias_validos) == 1: texto_data_hora = f", onde às {horas_validas[0]}, do dia {dias_validos[0]},"
                elif len(dias_validos) > 1:
                    str_horas = ", ".join(horas_validas[:-1]) + f" e {horas_validas[-1]}"
                    str_dias = ", ".join(dias_validos[:-1]) + f" e {dias_validos[-1]}"
                    texto_data_hora = f", onde às {str_horas}, dos dias {str_dias},"
                
                txt_endereco = f"à {endereco}" if endereco else "ao endereço informado no mesmo"
                txt_pessoa = f" a pessoa, Sr(a). {pessoa}" if pessoa else "a pessoa referida no mandado"
                
                txt_situacao = ""
                if situacao_simples == "Local Fechado": txt_situacao = "porque o local foi encontrado fechado e, mesmo após chamar várias vezes, ninguém atendeu."
                elif situacao_simples == "Pessoa Não Encontrada": txt_situacao = "porque não a encontrei no local."
                elif situacao_simples == "Não Localizei a Pessoa": txt_situacao = "porque não a localizei."
                
                paragrafo_unico = f"Certifico e dou fé que, em cumprimento ao mandado anexo, dirigi-me {txt_endereco}{texto_data_hora.rstrip(',')} e deixei de citar/intimar/notificar{txt_pessoa}, {txt_situacao} "
                
                if obteve_inf_simples == "Sim": paragrafo_unico += f"Conforme informações obtidas no local com o(a) Sr.(a) {nome_inf_simples}, este(a) informou que "
                elif obteve_inf_simples == "Não": paragrafo_unico += "Procurei obter informações junto aos moradores e vizinhos, não obtendo êxito, uma vez que ninguém forneceu informações. "
                elif obteve_inf_simples == "NQI": paragrafo_unico += "Conforme informações prestadas por um vizinho(a) que não quis se identificar, este(a) afirmou que "
                
                if obteve_inf_simples in ["Sim", "NQI"]:
                    infos = []
                    if motivo_simples == "Mudou-se": infos.append("a pessoa procurada não reside mais no local")
                    elif motivo_simples == "Não Reside": infos.append("a pessoa procurada não reside no local referido")
                    elif motivo_simples == "Falecido": infos.append("a pessoa procurada já se encontra falecida")
                    if nao_sabe_simples == "Não Conhece": infos.append("não a conhece")
                    if paradeiro_simples == "Incerto e Não Sabido": infos.append("encontra-se em local incerto e não sabido")
                    if infos:
                        texto_infos = "; ".join(infos[:-1]) + ", e que " + infos[-1] if len(infos) > 1 else infos[0]
                        paragrafo_unico += f"{texto_infos}. "
                        
                if condicao_simples == "Chuva": paragrafo_unico += "Certifico que a execução restou dificultada em virtude de intensa chuva. "
                if observacoes_simples: paragrafo_unico += f"{observacoes_simples} "
                
                doc = Document(); style = doc.styles['Normal']; font = style.font; font.name = 'Times New Roman'; font.size = Pt(12)
                try:
                    cabecalho_bytes = supabase.storage.from_("imagens_sistema").download("cabecalho.png")
                    p_img_cabecalho = doc.add_paragraph(); p_img_cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_cabecalho.add_run().add_picture(BytesIO(cabecalho_bytes), width=Cm(16))
                except: pass
                if processo:
                    texto_processo = f"Processo: {processo}"
                    if ano: texto_processo += f".{ano or '2026'}.8.13.{comarca}"
                    doc.add_paragraph(texto_processo)
                if mandado: doc.add_paragraph(f"Mandado nº: {mandado}")
                doc.add_paragraph(""); p_titulo = doc.add_paragraph(); run_titulo = p_titulo.add_run("CERTIDÃO"); run_titulo.bold = True; run_titulo.font.size = Pt(16); p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                doc.add_paragraph(paragrafo_unico.strip()).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; doc.paragraphs[-1].paragraph_format.first_line_indent = Pt(35.4); doc.add_paragraph("")
                doc.add_paragraph("Devolvo o mandado para os devidos fins. É verdade. Dou fé.").alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
                doc.add_paragraph(f"Santa Luzia, {data_certidao.day} de {meses[data_certidao.month - 1]} de {data_certidao.year}.").alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                
                try:
                    assinatura_bytes = supabase.storage.from_("assinaturas_usuarios").download(f"{usuario_atual}.png")
                    p_img_assinatura = doc.add_paragraph(); p_img_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_assinatura.add_run().add_picture(BytesIO(assinatura_bytes), width=Cm(6))
                except: pass 
                p_assinatura = doc.add_paragraph(); p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_nome = p_assinatura.add_run(f"{dados_usuario['nome']}\n"); run_nome.bold = True; run_nome.font.size = Pt(8)
                run_cargo = p_assinatura.add_run(f"{dados_usuario['cargo']}\n"); run_cargo.font.size = Pt(8)
                run_matricula = p_assinatura.add_run(f"{dados_usuario['matricula']}"); run_matricula.font.size = Pt(8)
                
                buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
                docx_bytes = buffer.getvalue()
                
                data_arquivo = hoje_real.strftime("%d-%m-%Y_%Hh%M")
                nome_base = f"Certidao_Simples_{processo}_{data_arquivo}"
                
                if formato_saida == "PDF (.pdf)":
                    arquivo_final_bytes = converter_docx_para_pdf(docx_bytes)
                    nome_final = nome_base + ".pdf"
                    mime_final = "application/pdf"
                    if not arquivo_final_bytes:
                        st.error("Erro na conversão PDF. Baixando DOCX.")
                        arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                supabase.storage.from_("certidoes_usuarios").upload(file=arquivo_final_bytes, path=f"{usuario_atual}/{nome_final}", file_options={"content-type": mime_final})
                
                st.session_state['doc_pronto_bytes_b'] = arquivo_final_bytes
                st.session_state['doc_pronto_nome_b'] = nome_final
                st.session_state['doc_pronto_mime_b'] = mime_final
                st.session_state['piscar_tela'] = True
                st.session_state['limpar_simples'] = True
                st.rerun()

        if 'doc_pronto_bytes_b' in st.session_state:
            if st.session_state.get('piscar_tela'):
                st.balloons(); st.toast("✅ Certidão gerada!", icon="🎉")
                st.session_state['piscar_tela'] = False
            st.success("✅ Certidão simples salva na Nuvem!")
            st.download_button("📥 Baixar Arquivo", data=st.session_state['doc_pronto_bytes_b'], file_name=st.session_state['doc_pronto_nome_b'], mime=st.session_state['doc_pronto_mime_b'], type="primary", use_container_width=True)

    # ==========================================
    # OPÇÃO C: CERTIDÃO POSITIVA
    # ==========================================
    elif tipo_certidao == "Certidão Positiva":
        
        if st.session_state.get('limpar_positiva'):
            st.session_state["fin_pos"] = "Citação"
            st.session_state["mod_pos"] = "Presencial"
            st.session_state["contra_pos"] = "Sim"
            st.session_state["ass_pos"] = "Sim"
            st.session_state["adv_pos"] = "Não Perguntado"
            st.session_state["obs_pos"] = ""
            st.session_state['limpar_positiva'] = False
            
        st.subheader("Detalhes da Diligência Positiva")
        finalidade_pos = st.selectbox("Finalidade principal do Mandado:", ["Citação", "Intimação", "Notificação", "Penhora", "Avaliação"], key="fin_pos")
        st.markdown("---")
        
        c_mod, c_contra, c_ass = st.columns(3)
        with c_mod: mod_pos = st.radio("Como foi o contato?", ["Presencial", "Telefone/WhatsApp"], key="mod_pos")
        with c_contra: contrafe_pos = st.radio("Aceitou a contrafé?", ["Sim", "Não"], key="contra_pos")
        with c_ass: ass_pos = st.radio("Colheu assinatura?", ["Sim", "Não", "Covid-19"], key="ass_pos")
            
        st.markdown("---")
        c_adv, c_obs = st.columns([1, 2])
        with c_adv: adv_pos = st.radio("Perguntou sobre advogado?", ["Não Perguntado", "Tem condições", "Não tem condições"], key="adv_pos")
        with c_obs: obs_pos = st.text_area("Observações Adicionais", height=110, key="obs_pos")
            
        st.divider()
        
        if st.button("Salvar na Nuvem / Gerar Documento", type="primary", use_container_width=True, key="btn_gerar_positiva"):
            with st.spinner("Gerando certidão positiva..."):
                dias_validos = [d for d in [d1, d2, d3] if d]
                horas_cruas = [h for h in [h1, h2, h3] if h]
                horas_validas = []
                for h in horas_cruas:
                    h_limpo = h.strip()
                    if h_limpo and not h_limpo.lower().endswith(('h', 'hs')): h_limpo += 'h'
                    horas_validas.append(h_limpo)
                
                dia_pos = dias_validos[-1] if dias_validos else "___/___"
                hora_pos = horas_validas[-1] if horas_validas else "___:___"

                verbo_acao = ""
                if finalidade_pos == "Citação": verbo_acao = "à CITAÇÃO"
                elif finalidade_pos == "Intimação": verbo_acao = "à INTIMAÇÃO"
                elif finalidade_pos == "Notificação": verbo_acao = "à NOTIFICAÇÃO"
                elif finalidade_pos == "Penhora": verbo_acao = "à PENHORA"
                elif finalidade_pos == "Avaliação": verbo_acao = "à AVALIAÇÃO"

                txt_pessoa = f" de {pessoa}" if pessoa else " da pessoa referida no mandado"
                txt_endereco = f"à {endereco}" if endereco else "ao endereço indicado"

                if mod_pos == "Telefone/WhatsApp":
                    paragrafo = f"Certifico e dou fé que, em cumprimento ao mandado anexo, no dia {dia_pos}, por volta das {hora_pos}, procedi, por via remota (ligação telefônica/aplicativo de mensagens), {verbo_acao}{txt_pessoa}. "
                    if contrafe_pos == "Sim": paragrafo += "Na oportunidade, encaminhei a respectiva contrafé eletronicamente, a qual teve seu recebimento confirmado. "
                    else: paragrafo += "Tentei encaminhar a respectiva contrafé eletronicamente, porém a mesma foi recusada ou não teve seu recebimento confirmado. "
                else:
                    paragrafo = f"Certifico e dou fé que, em cumprimento ao mandado anexo, dirigi-me {txt_endereco}, ocasião em que, no dia {dia_pos}, por volta das {hora_pos}, procedi {verbo_acao}{txt_pessoa}. "
                    if contrafe_pos == "Sim": paragrafo += "Na oportunidade, li-lhe o mandado e entreguei-lhe a respectiva contrafé, a qual foi aceita pela parte"
                    else: paragrafo += "Na oportunidade, li-lhe o mandado e ofereci-lhe a respectiva contrafé, a qual foi recusada pela parte"
                        
                    if ass_pos == "Sim": paragrafo += ", que exarou sua assinatura no documento. "
                    elif ass_pos == "Não": paragrafo += ", que se recusou a exarar sua assinatura no documento. "
                    else: paragrafo += ", deixando eu de colher a assinatura física como medida de prevenção sanitária/Covid-19. "

                if adv_pos == "Tem condições": paragrafo += "Questionado(a), a parte informou possuir condições de constituir um advogado particular para sua defesa. "
                elif adv_pos == "Não tem condições": paragrafo += "Questionado(a), a parte declarou ser hipossuficiente, necessitando da nomeação de um defensor para atuar em sua defesa. "
                if obs_pos: paragrafo += f"{obs_pos.strip()} "

                doc = Document(); style = doc.styles['Normal']; font = style.font; font.name = 'Times New Roman'; font.size = Pt(12)
                try:
                    cabecalho_bytes = supabase.storage.from_("imagens_sistema").download("cabecalho.png")
                    p_img_cabecalho = doc.add_paragraph(); p_img_cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_cabecalho.add_run().add_picture(BytesIO(cabecalho_bytes), width=Cm(16))
                except: pass
                if processo:
                    texto_processo = f"Processo: {processo}"
                    if ano: texto_processo += f".{ano or '2026'}.8.13.{comarca}"
                    doc.add_paragraph(texto_processo)
                if mandado: doc.add_paragraph(f"Mandado nº: {mandado}")
                doc.add_paragraph(""); p_titulo = doc.add_paragraph(); run_titulo = p_titulo.add_run("CERTIDÃO POSITIVA"); run_titulo.bold = True; run_titulo.font.size = Pt(16); p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                doc.add_paragraph(paragrafo.strip()).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; doc.paragraphs[-1].paragraph_format.first_line_indent = Pt(35.4); doc.add_paragraph("")
                doc.add_paragraph("Devolvo o mandado para os devidos fins. É verdade. Dou fé.").alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
                doc.add_paragraph(f"Santa Luzia, {data_certidao.day} de {meses[data_certidao.month - 1]} de {data_certidao.year}.").alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                try:
                    assinatura_bytes = supabase.storage.from_("assinaturas_usuarios").download(f"{usuario_atual}.png")
                    p_img_assinatura = doc.add_paragraph(); p_img_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_assinatura.add_run().add_picture(BytesIO(assinatura_bytes), width=Cm(6))
                except: pass 
                p_assinatura = doc.add_paragraph(); p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_nome = p_assinatura.add_run(f"{dados_usuario['nome']}\n"); run_nome.bold = True; run_nome.font.size = Pt(8)
                run_cargo = p_assinatura.add_run(f"{dados_usuario['cargo']}\n"); run_cargo.font.size = Pt(8)
                run_matricula = p_assinatura.add_run(f"{dados_usuario['matricula']}"); run_matricula.font.size = Pt(8)
                
                buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
                docx_bytes = buffer.getvalue()
                
                data_arquivo = hoje_real.strftime("%d-%m-%Y_%Hh%M")
                nome_base = f"Certidao_Positiva_{processo}_{data_arquivo}"
                
                if formato_saida == "PDF (.pdf)":
                    arquivo_final_bytes = converter_docx_para_pdf(docx_bytes)
                    nome_final = nome_base + ".pdf"
                    mime_final = "application/pdf"
                    if not arquivo_final_bytes:
                        st.error("Erro na conversão PDF. Baixando DOCX.")
                        arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                supabase.storage.from_("certidoes_usuarios").upload(file=arquivo_final_bytes, path=f"{usuario_atual}/{nome_final}", file_options={"content-type": mime_final})
                
                st.session_state['doc_pronto_bytes_c'] = arquivo_final_bytes
                st.session_state['doc_pronto_nome_c'] = nome_final
                st.session_state['doc_pronto_mime_c'] = mime_final
                st.session_state['piscar_tela'] = True
                st.session_state['limpar_positiva'] = True
                st.rerun()

        if 'doc_pronto_bytes_c' in st.session_state:
            if st.session_state.get('piscar_tela'):
                st.balloons(); st.toast("✅ Certidão Positiva gerada!", icon="🎉")
                st.session_state['piscar_tela'] = False
            st.success("✅ Certidão salva na Nuvem!")
            st.download_button("📥 Baixar Arquivo", data=st.session_state['doc_pronto_bytes_c'], file_name=st.session_state['doc_pronto_nome_c'], mime=st.session_state['doc_pronto_mime_c'], type="primary", use_container_width=True)
            
    # ==========================================
    # OPÇÃO D: CERTIDÃO POSITIVA POR HORA CERTA
    # ==========================================
    elif tipo_certidao == "Certidão Positiva por Hora Certa":
        
        if st.session_state.get('limpar_horacerta'):
            st.session_state["fin_hc"] = "Citação"
            for k in ["hc_nome_terceiro", "hc_relacao", "hc_data_retorno", "hc_hora_retorno"]:
                st.session_state[k] = ""
            st.session_state["hc_encontrou_alvo"] = "Não"
            st.session_state["hc_aceitou"] = "Sim"
            st.session_state["hc_assinou"] = "Não"
            st.session_state['limpar_horacerta'] = False
            
        st.subheader("1. Suspeita de Ocultação e Agendamento")
        finalidade_hc = st.selectbox("Ato sendo praticado:", ["Citação", "Intimação", "Notificação"], key="fin_hc")
        
        c_hc1, c_hc2 = st.columns(2)
        with c_hc1:
            hc_nome_terceiro = st.text_input("Com quem marcou a hora certa?", placeholder="Ex: Sra. Teresinha dos Santos", key="hc_nome_terceiro")
            hc_data_retorno = st.text_input("Data marcada para o retorno:", placeholder="Ex: 19/05", key="hc_data_retorno")
        with c_hc2:
            hc_relacao = st.text_input("Qualidade/Relação:", placeholder="Ex: mãe, vizinha, porteiro", key="hc_relacao")
            hc_hora_retorno = st.text_input("Hora marcada para o retorno:", placeholder="Ex: 18:15", key="hc_hora_retorno")

        st.markdown("---")
        st.subheader("2. Desfecho do Retorno")
        c_hc3, c_hc4, c_hc5 = st.columns(3)
        with c_hc3: hc_encontrou_alvo = st.radio("Encontrou a pessoa?", ["Não", "Sim"], key="hc_encontrou_alvo")
        with c_hc4: hc_aceitou = st.radio("O terceiro/alvo aceitou receber?", ["Sim", "Não"], key="hc_aceitou")
        with c_hc5: hc_assinou = st.radio("O terceiro/alvo assinou?", ["Não", "Sim", "Covid-19"], key="hc_assinou")

        st.divider()

        if st.button("Salvar na Nuvem / Gerar Documento", type="primary", use_container_width=True, key="btn_gerar_horacerta"):
            with st.spinner("Construindo certidão de Hora Certa..."):
                
                txt_endereco = f"à {endereco}" if endereco else "ao endereço indicado"
                txt_pessoa = f" de {pessoa}" if pessoa else " da pessoa referida no mandado"
                
                dias_validos = [d for d in [d1, d2, d3] if d]
                horas_cruas = [h for h in [h1, h2, h3] if h]
                horas_validas = []
                for h in horas_cruas:
                    h_limpo = h.strip()
                    if h_limpo and not h_limpo.lower().endswith(('h', 'hs', 'min')): h_limpo += 'h'
                    horas_validas.append(h_limpo)
                
                texto_data_hora = ""
                if len(dias_validos) == 1: texto_data_hora = f"no dia {dias_validos[0]}, por volta das {horas_validas[0]}, ocasião em que"
                elif len(dias_validos) > 1:
                    str_horas = ", ".join(horas_validas[:-1]) + f" e {horas_validas[-1]}"
                    str_dias = ", ".join(dias_validos[:-1]) + f" e {dias_validos[-1]}"
                    texto_data_hora = f"nos dias {str_dias}, por volta das {str_horas}, respectivamente, ocasiões em que"
                
                hr_limpo = hc_hora_retorno.strip()
                if hr_limpo and not hr_limpo.lower().endswith(('h', 'hs', 'min')): hr_limpo += 'h'

                nome_ato = finalidade_hc.upper()

                txt_terceiro = f"na pessoa do(a) Sr(a). {hc_nome_terceiro}" if hc_nome_terceiro else "na pessoa de um terceiro ali presente"
                txt_relacao = f", na qualidade de {hc_relacao}," if hc_relacao else ","
                txt_retorno_alvo = "ali não a encontrando" if hc_encontrou_alvo == "Não" else "ali a encontrando"
                
                if hc_aceitou == "Sim":
                    if hc_assinou == "Sim": txt_final = "a qual aceitou o documento e exarou sua assinatura no respectivo mandado."
                    elif hc_assinou == "Não": txt_final = "a qual aceitou o documento, mas recusou-se a exarar sua assinatura no respectivo mandado."
                    else: txt_final = "a qual aceitou o documento, deixando eu de colher a assinatura física como medida de prevenção sanitária/Covid-19."
                else: txt_final = "a qual se recusou a receber a contrafé e a assinar o respectivo mandado."

                paragrafo = f"Certifico e dou fé que, em cumprimento ao mandado anexo, dirigi-me {txt_endereco}, onde {texto_data_hora} não encontrei a pessoa procurada. Diante das diligências frustradas e havendo fundada suspeita de ocultação, efetuei o agendamento de HORA CERTA {txt_terceiro}{txt_relacao} intimando-o(a) de que retornaria no dia {hc_data_retorno}, pontualmente às {hr_limpo}, para efetivar o ato judicial. Retornando no dia e hora estritamente designados, {txt_retorno_alvo}, dei por realizada a {nome_ato}{txt_pessoa}, deixando a respectiva contrafé com a pessoa mencionada, {txt_final}"
                
                doc = Document(); style = doc.styles['Normal']; font = style.font; font.name = 'Times New Roman'; font.size = Pt(12)
                try:
                    cabecalho_bytes = supabase.storage.from_("imagens_sistema").download("cabecalho.png")
                    p_img_cabecalho = doc.add_paragraph(); p_img_cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_cabecalho.add_run().add_picture(BytesIO(cabecalho_bytes), width=Cm(16))
                except: pass
                if processo:
                    texto_processo = f"Processo: {processo}"
                    if ano: texto_processo += f".{ano or '2026'}.8.13.{comarca}"
                    doc.add_paragraph(texto_processo)
                if mandado: doc.add_paragraph(f"Mandado nº: {mandado}")
                doc.add_paragraph(""); p_titulo = doc.add_paragraph(); run_titulo = p_titulo.add_run("CERTIDÃO POSITIVA POR HORA CERTA"); run_titulo.bold = True; run_titulo.font.size = Pt(16); p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                doc.add_paragraph(paragrafo.strip()).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; doc.paragraphs[-1].paragraph_format.first_line_indent = Pt(35.4); doc.add_paragraph("")
                doc.add_paragraph("Devolvo o mandado para os devidos fins. É verdade. Dou fé.").alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
                doc.add_paragraph(f"Santa Luzia, {data_certidao.day} de {meses[data_certidao.month - 1]} de {data_certidao.year}.").alignment = WD_ALIGN_PARAGRAPH.CENTER; doc.add_paragraph("")
                try:
                    assinatura_bytes = supabase.storage.from_("assinaturas_usuarios").download(f"{usuario_atual}.png")
                    p_img_assinatura = doc.add_paragraph(); p_img_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img_assinatura.add_run().add_picture(BytesIO(assinatura_bytes), width=Cm(6))
                except: pass 
                p_assinatura = doc.add_paragraph(); p_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_nome = p_assinatura.add_run(f"{dados_usuario['nome']}\n"); run_nome.bold = True; run_nome.font.size = Pt(8)
                run_cargo = p_assinatura.add_run(f"{dados_usuario['cargo']}\n"); run_cargo.font.size = Pt(8)
                run_matricula = p_assinatura.add_run(f"{dados_usuario['matricula']}"); run_matricula.font.size = Pt(8)
                
                buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
                docx_bytes = buffer.getvalue()
                
                data_arquivo = hoje_real.strftime("%d-%m-%Y_%Hh%M")
                nome_base = f"Certidao_HoraCerta_{processo}_{data_arquivo}"
                
                if formato_saida == "PDF (.pdf)":
                    arquivo_final_bytes = converter_docx_para_pdf(docx_bytes)
                    nome_final = nome_base + ".pdf"
                    mime_final = "application/pdf"
                    if not arquivo_final_bytes:
                        st.error("Erro na conversão PDF. Baixando DOCX.")
                        arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    arquivo_final_bytes, nome_final, mime_final = docx_bytes, nome_base + ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                supabase.storage.from_("certidoes_usuarios").upload(file=arquivo_final_bytes, path=f"{usuario_atual}/{nome_final}", file_options={"content-type": mime_final})
                
                st.session_state['doc_pronto_bytes_d'] = arquivo_final_bytes
                st.session_state['doc_pronto_nome_d'] = nome_final
                st.session_state['doc_pronto_mime_d'] = mime_final
                st.session_state['piscar_tela'] = True
                st.session_state['limpar_horacerta'] = True 
                st.rerun() 

        if 'doc_pronto_bytes_d' in st.session_state:
            if st.session_state.get('piscar_tela'):
                st.balloons(); st.toast("✅ Certidão de Hora Certa gerada!", icon="🎉")
                st.session_state['piscar_tela'] = False 
            st.success("✅ Certidão salva na Nuvem!")
            st.download_button("📥 Baixar Arquivo", data=st.session_state['doc_pronto_bytes_d'], file_name=st.session_state['doc_pronto_nome_d'], mime=st.session_state['doc_pronto_mime_d'], type="primary", use_container_width=True)
